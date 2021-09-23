#import required modules
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx2pdf import convert

#-----------------------------------------------GENERATE BASELINE ENERGY REPORT----------------------------------------------------
def energyBaseline(path):
    print('Generating report for Baseline Energy function...')

    document = Document()

    document.add_heading('Baseline Energy - Analysis Report', 0) #Report title

    #Introductory paragraph
    p = document.add_paragraph('The baseline energy function ')
    p.add_run('compares energy use during operating hours (workhours) and outside operating hours (after-hours) for heating, cooling, and electricity*. ').bold=True
    p.add_run("This function is intended to help the user assess the effectiveness of schedules and their ability to reduce \
energy use outside of a building's operating hours. Plots are generated which compare the rate of energy use during and outside operating \
hours with respect to outdoor air temperature, and predict energy consumption at certain outdoor air temperatures - these are \
done separately for heating, cooling, and electricity*. The generated key performance indicators (KPI) quantify schedule \
effectiveness and afterhours energy use. More information is found in the respective sections.")
    p = document.add_paragraph("*Electricity use comparison is shown only if electricity use analysis was conducted.")

    #Visualization heading and description
    document.add_heading('Visuals', level=1)
    p = document.add_paragraph('The first set of visuals (to the left) compare the rate of energy use during operating hours (workhours) and \
outside operating hours (after-hours) as a function of outdoor air temperatures - this is done separately for heating, \
cooling, and electricity. Current schedules may be ineffective at reducing after-hours energy use if the inclined slope of the after-hours energy use \
line is similiar or identical to the workhours energy use line. If the slope of the after-hours energy use line is shallower \
than the workhour energy use line, the current schedules are effective at reducing energy use outside operating hours.')
    p = document.add_paragraph('The second set of visuals (to the right) illustrates the sensitivity of energy use with respect to outdoor \
air temperatures - this is done separately for heating, cooling, and electricity. If the lines are spaced considerably apart, \
the energy use is particularily sensitive to fluctuations in outdoor air temperature.')

    #Add visualizations
    document.add_picture(os.path.join(path,'energyBase_heating.png'), width=Inches(5.75))
    document.add_picture(os.path.join(path,'energyBase_cooling.png'), width=Inches(5.75))

    try:
        document.add_picture(os.path.join(path,'energyBase_electricity.png'), width=Inches(5.75))
        os.remove(os.path.join(path,'energyBase_electricity.png'))
    except:
        pass

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('The key performance indicators (KPIs) are ')
    p.add_run('Schedule effectiveness').bold = True
    p.add_run(' and ')
    p.add_run('Afterhours energy use ratio').bold = True
    p.add_run('. Schedule effectiveness quantifies the difference between the slope of the workhours energy use line and the \
afterhours energy use line. Values approaching zero (0%) indicate similiar or identical inclined slopes, positive (+) values indicate \
a steeper workhours energy use slope than afterhours, and negative (-) values indicate a steeper afterhours energy use slope. Therefore, \
a greater positive value is desirable since it indicates an effective reduction in energy use rates during afterhours.\n\n\
The Afterhours energy use ratio is the ratio of energy use during afterhours over the total energy use. A higher value indicates a larger \
portion of the total energy consumption used during after-hours. Therefore, a lower value is desirable.')

    #Extract data from excel sheet and add table to document
    kpis = pd.read_excel(os.path.join(path,'energyBase_summary.xlsx'),sheet_name='KPIs')
    kpis.drop(kpis.columns[0],axis=1,inplace=True)
    kpis.iloc[:,1:] = (kpis.iloc[:,1:] * 100).astype(int).astype(str) + '%'

    p = document.add_heading('Schedule Effectiveness and Afterhours energy use ratio',level=2)
    t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
    
    for j in range(kpis.shape[-1]): # add the header rows.
        t.cell(0,j).text = kpis.columns[j]
    
    for i in range(kpis.shape[0]): # add the rest of the data frame
        for j in range(kpis.shape[-1]):
            t.cell(i+1,j).text = str(kpis.values[i,j])
    
    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(os.path.join(path,'energyBaseline_report.docx'))
    convert(os.path.join(path,'energyBaseline_report.docx'), os.path.join(path,'report.pdf'))

    #Remove all used files
    os.remove(os.path.join(path,'energyBase_heating.png'))
    os.remove(os.path.join(path,'energyBase_cooling.png'))
    os.remove(os.path.join(path,'energyBase_summary.xlsx'))
    os.remove(os.path.join(path,'energyBaseline_report.docx'))

    print('Report successfully generated!')

    return

#-----------------------------------------------GENERATE AHU ANOMALY REPORT----------------------------------------------------
def ahuAnomaly(path):
    print('Generating report for AHU anomaly detection function...')

    #Extract KPIs excel sheet
    faults_df = pd.read_excel(os.path.join(path,'ahu_anomaly_summary.xlsx'),sheet_name='faults')
    faults_df.drop(faults_df.columns[0],axis=1,inplace=True)
    faults_df.iloc[:,1] = (faults_df.iloc[:,1]).astype(int).astype(str) + '%'

    document = Document()

    document.add_heading('AHU Anomaly - Analysis Report', 0) #Report title

    #Introductory paragraph
    p = document.add_paragraph('The AHU anomaly detection function ')
    p.add_run('detects hard and soft faults related to air handling units (AHUs). ').bold=True
    p.add_run("This function helps identify potential causes for anomalous AHU operations which may cause \
energy use inefficiencies. The visuals are intended to aid understanding of the detected faults. These depict \
supply air temperature, and the coolest/warmest/average return air temperatures as a function of outdoor air temperature, and \
damper and valve actuator positions as a function of outdoor air temperature. Additionally, a number of diagrams are generated which \
depict damper and valve actuator positions and temperature readings at characteristic AHU operational periods. More information is \
available at the respective sections.")

    #Visualization heading and description - Part 1
    document.add_heading('Visuals - Split-range controller', level=1)
    p = document.add_paragraph('A set of two charts are generated for each AHU dataset inputted. The first (top) plots supply air \
temperature, and the coolest/warmest/average return air temperatures as a function of outdoor air temperature. For \
reference, the "ideal" supply air temperature is depicted. The second (bottom) chart is a Split-range controller diagram, which \
plots the outdoor air damper position (OA), heating coil valve position (HC), cooling coil valve position (CC) and average fraction \
of active perimeter heaters (RAD) with respect to outdoor air temperature. The four underlaying color zones represent the four \
distinct operating mode: Heating (red zone), economizer (yellow zone), economizer with cooling (grey zone), and cooling (blue zone). \
As an example, the below Split-range controller diagram is representative of normal, healthy AHU operations. Some ')
    p.add_run('key characteristics of normal, healthy AHU operation').bold=True
    p.add_run(' include:')

    document.add_paragraph('Heating coil active ONLY in heating mode', style='List Bullet')
    document.add_paragraph('Cooling coil active ONLY in economizer with cooling and cooling mode', style='List Bullet')
    document.add_paragraph('Heating and cooling coils should not operate simutaneously', style='List Bullet')
    document.add_paragraph('Perimeter heating should be minimal in economizer mode. This can be achieved by increasing the supply air \
temperature setpoint in the heating season, while monitoring any occurence of overheating.', style='List Bullet')

    cwd = os.getcwd()
    document.add_picture(os.path.join(cwd,'static','img','SRC_example.png'), width=Inches(5.75))

    p = document.add_paragraph("Suboptimal supply air temperatures can result excessive energy consumption from excessive perimeter heating use, \
economizing, or fan use. To guide supply air temperature setpoint adjustments, the typical 'ideal' supply air temperature range is provided as a \
reference.\n\nFirst, ensure that the supply air temperature is not constant and is controlled by a program. Review the BAS program responsible for \
the supply air temperature setpoint analog variable.\n\nIf the supply air temperature is ABOVE the 'ideal' range in the cooling season, excessive \
fan power to deliver required necessary cooling may result. Inspect what may be causing this in the BAS program. Likely causes include:")
    document.add_paragraph('Inappropriate logic in the BAS program', style='List Bullet')
    document.add_paragraph('VAV terminal unit dampers stuck open causing overcooling in select rooms', style='List Bullet')
    document.add_paragraph("If the supply air temperature is BELOW the 'ideal' range in the heating season, excessive use of the perimeter heaters \
or economizing may result. Likely causes include:")
    document.add_paragraph('Inappropriate logic in the BAS program', style='List Bullet')
    document.add_paragraph('VAV terminal unit dampers stuck closed', style='List Bullet')
    document.add_paragraph('Radiators stuck on', style='List Bullet')
    document.add_paragraph('Overheating rooms caused by internal/external heat gains', style='List Bullet')
    document.add_paragraph("In both of these cases, ensure that the airflow and temperature sensors work \
as intended in these rooms.")

    #For each AHU, output the first set of visuals
    for i in faults_df.index:

        document.add_heading('AHU: ' + str(faults_df.iloc[i][0]), level=2)
        document.add_picture(os.path.join(path,'f2a_ahu_'+str(i+1)+'.png'), width=Inches(5))
        os.remove(os.path.join(path,'f2a_ahu_'+str(i+1)+'.png'))
    
    #Visualization heading and description - Part 2
    document.add_heading('Visuals - AHU operating periods', level=1)
    p = document.add_paragraph("A set of four to six visuals per AHU are generated which depict characteristic operating \
periods of the AHU and the average damper and valve positions and temperatures at those periods. The fraction of time of \
operation is the percentage of the total time of the AHU's operation which exhibit the displayed damper/valve positions \
and temperatures.")
    
    #For each AHU, output the second set of visuals
    for i in faults_df.index:

        document.add_heading('AHU: ' + str(faults_df.iloc[i][0]), level=2)

        for j in range(1,7):
            try:
                document.add_picture(os.path.join(path,'f2b_ahu_'+str(i+1)+'_C_'+str(j)+'.png'), width=Inches(5))
                os.remove(os.path.join(path,'f2b_ahu_'+str(i+1)+'_C_'+str(j)+'.png'))
            except FileNotFoundError:
                break

    document.add_page_break()

    #Extract data from excel sheet and add table to document

    p = document.add_heading('Key performance indicators - AHU faults',level=1)
    p = document.add_paragraph('The following table lists the hard and soft faults identified by the function. The AHU health \
index is also provided for each AHU which is '+ r'100%'+' if no faults are detected and '+ r'0%'+' if all six faults are detected. \
The following faults may be detected:')

    document.add_paragraph('Cooling coil stuck: This fault is generated if no or minimal use of the cooling coil was observed in the data. \
This may be symptomatic of a faulty sensor, a stuck valve, or a conflict in the operational logic.', style='List Bullet')
    document.add_paragraph('Heating coil stuck: This fault is generated if no or minimal use of the cooling coil was observed in the data. \
This may be symptomatic of a faulty sensor, a stuck valve, or a conflict in the operational logic.', style='List Bullet')
    document.add_paragraph("Check economizer logic: This fault is generated if excessive use of the perimeter heatings was observed in the \
AHU's economizer mode.", style='List Bullet')
    document.add_paragraph('Low/High outdoor air: This fault is generated if an inadequate or excessive amount of outdoor air was observed. \
This may be symptomatic of a stuck outdoor air damper or faulty damper sensor/actuator.', style='List Bullet')
    document.add_paragraph('Check mode of operation logic: This fault is generated if the weekly operational time (i.e. when the AHU fans are \
operational) exceeds 100 hours a week. This is typically considered excessive operations. It is suggested to check the operational logic for any conflicts \
which may result in unintended operation of the AHUs', style='List Bullet')
    document.add_paragraph("Check supply air temperature reset logic: This fault is generated if a setpoint reset scheme was not detected or there \
is a logic mistake in the reset program. Ensuring that an appropriate setpoint reset scheme is present and operating as intended, there may be excessive use of the perimeter heating \
devices during the AHUs' economizer mode. This may be a result of select rooms overheating in the heating season. If this is the case, consider \
increasing the maximum terminal airflow setpoints in these overheating rooms.", style='List Bullet')

    t = document.add_table(faults_df.shape[0]+1, faults_df.shape[1])
    
    for j in range(faults_df.shape[-1]): # add the header rows.
        t.cell(0,j).text = faults_df.columns[j]
    
    for i in range(faults_df.shape[0]): # add the rest of the data frame
        for j in range(faults_df.shape[-1]):
            t.cell(i+1,j).text = str(faults_df.values[i,j])

    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(os.path.join(path,'ahuAnomaly_report.docx'))
    convert(os.path.join(path,'ahuAnomaly_report.docx'), os.path.join(path,'report.pdf'))
    print('Report successfully generated!')

    #Remove used files
    os.remove(os.path.join(path,'ahuAnomaly_report.docx'))
    os.remove(os.path.join(path,'ahu_anomaly_summary.xlsx'))

    return

#-----------------------------------------------GENERATE ZONE ANOMALY REPORT----------------------------------------------------
def zoneAnomaly(path):
    print('Generating report for zone anomaly detection function...')

    #Extract KPIs excel sheet
    kpis_heating = pd.read_excel(os.path.join(path,'zone_anomaly_summary.xlsx'),sheet_name='Htg_summary')
    kpis_heating.drop(kpis_heating.columns[0],axis=1,inplace=True)

    kpis_cooling = pd.read_excel(os.path.join(path,'zone_anomaly_summary.xlsx'),sheet_name='Clg_summary')
    kpis_cooling.drop(kpis_cooling.columns[0],axis=1,inplace=True)

    samples_heating = pd.read_excel(os.path.join(path,'zone_anomaly_summary.xlsx'),sheet_name='Htg_samples',keep_default_na=False)
    samples_heating.drop(samples_heating.columns[0],axis=1,inplace=True)

    samples_cooling = pd.read_excel(os.path.join(path,'zone_anomaly_summary.xlsx'),sheet_name='Clg_samples',keep_default_na=False)
    samples_cooling.drop(samples_cooling.columns[0],axis=1,inplace=True)

    #Generate Word document
    document = Document()

    #Report title
    document.add_heading('Zone Anomaly - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The zone anomaly function ')
    p.add_run('detects anomalous zones based on indoor air temperature and airflow control errors. ').bold = True
    p.add_run('This function can help identify potential faults in variable air volume (VAV) units which \
may result in anomalous airflow and temperature conditions in thermal zones. Visuals are also generated which \
plot the average indoor air temperature and airflow control error of groups of zones; this is done separately \
for the heating and cooling season.')

    #Visualization heading and description
    document.add_heading('Visuals', level=1)
    p = document.add_paragraph('The generated visuals plot the average indoor air temperature and average \
airflow control error of each group of zones. Each group of zones (C1, C2, C3, etc.) represents a number of zones \
which is presented at the top-left of the visual. The airflow control error is the difference between the \
actual and setpoint airflow rate with respect to the airflow rate setpoint. A positive (+) percentage indicates a \
higher actual flowrate than the setpoint and a negative (-) percentage indicates a lower actual flowrate than the \
setpoint. This visual is generated twice, once for the heating season (December through February) and again \
for the cooling season (May through August). The following are possible symptoms of anomalous operations:')

    document.add_paragraph('High flow: Zones exhibit abnormally high airflow. Ensure the VAV terminal damper and \
airflow sensors are operating as intended.', style='List Bullet')
    document.add_paragraph('Low flow: Zones exhibit abnormally low airflow. Ensure the VAV terminal damper and \
airflow sensors are operating as intended.', style='List Bullet')
    document.add_paragraph("Cold: Zones exhibit abnormally low indoor air temperatures. If zones exhibit acceptable \
airflow control errors (i.e., high/low flow faults are not present), ensure perimeter heating devices and/or reheat coils are operating as intended. Consider \
decreasing the VAV terminal minimum airflow setpoint.", style='List Bullet')
    document.add_paragraph("Hot: Zones exhibit abnormally high indoor air temperatures. If zones exhibit acceptable \
airflow control errors (i.e., high/low flow faults are not present), ensure perimeter heating devices and/or reheat coils are operating as intended. Consider \
increasing the VAV terminal maximum airflow setpoint.", style='List Bullet')

    #Add visualizations
    document.add_picture(os.path.join(path,'zone_heat_season.png'), width=Inches(4.75))
    document.add_picture(os.path.join(path,'zone_cool_season.png'), width=Inches(4.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('This section presents the generated KPIs - ')
    p.add_run('the zone health index').bold = True
    p.add_run('. The zone health index is the ') 
    p.add_run('ratio of zones with acceptable indoor air temperature and airflow control error over the \
total number of zones ').bold = True
    p.add_run('- this is calculated separately for the heating season (December through February) and the \
cooling season (May through August). An acceptable indoor air temperature is considered to be between 20 and \
25 degrees, and an acceptable airflow control error is +/- 20%. A greater percentage is desirable since this \
indicates little to no detected anomalous zones.')

    heating_health_index = (round(kpis_heating.iloc[0]['Health Index'],3))*100
    cooling_health_index = (round(kpis_cooling.iloc[0]['Health Index'],3))*100

    p = document.add_paragraph('Zone health index for heating season: ', style='List Bullet')
    p.add_run(str(heating_health_index) + '%').bold = True
    p = document.add_paragraph('Zone health index for cooling season: ', style='List Bullet')
    p.add_run(str(cooling_health_index) + '%').bold = True

    #Add tables with summary tables
    p = document.add_paragraph('The following tables are a summary of the clustering results, which lists the number of zones in each group and the average \
indoor air temperature and average airflow control error for each group. For the heating season, the average fraction \
of active (ON-state) perimeter heaters within a zone is also provided for each cluster.')

    kpis_heating.drop(kpis_heating.columns[-1],axis=1,inplace=True) #Drop the health index from the heating season table
    kpis_heating.iloc[:,0] = (kpis_heating.iloc[:,0]).astype(int)
    kpis_heating.iloc[:,2:] = (kpis_heating.iloc[:,2:]).round(1).astype(str) + '%'

    kpis_cooling.drop(kpis_cooling.columns[-1],axis=1,inplace=True) #Drop the health index from the cooling season table
    kpis_cooling.iloc[:,0] = (kpis_cooling.iloc[:,0]).astype(int)
    kpis_cooling.iloc[:,2] = (kpis_cooling.iloc[:,2]).round(1).astype(str) + '%'

    table_labels = ['Number of zones','Average indoor air temperature (C)','Average airflow control error','Average fraction of active perimeter heaters']
    table_headings = ['Cluster summary - Heating season (December-February)','Cluster summary - Cooling season (May-August)']
    k = 0

    for df in [kpis_heating,kpis_cooling]:
        df = df.round(1)
        p = document.add_heading(table_headings[k],level=2)
        t = document.add_table(df.shape[0]+1, df.shape[1])
        
        for j in range(df.shape[-1]): # add the header rows.
            t.cell(0,j).text = table_labels[j]
        
        for i in range(df.shape[0]): # add the rest of the data frame
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        k+=1
        t.style = 'Colorful List'
    
    document.add_page_break()
    
    #Add tables with cluster samples
    p = document.add_paragraph('The following tables lists the individual zones (by filename) by its resultant group (i.e., the zones listed under C1 \
belong to the zone cluster C1.) These tables can be used to identify the zone(s) which were determined to exhibit anomalous conditions.')

    table_headings = ['Zone samples - Heating season (December-February)','Zone samples - Cooling season (May-August)']
    k = 0

    for df in [samples_heating,samples_cooling]:
        p = document.add_heading(table_headings[k],level=2)
        t = document.add_table(df.shape[0]+1, df.shape[1])
        
        for j in range(df.shape[-1]): # add the header rows.
            t.cell(0,j).text = df.columns[j]
        
        for i in range(df.shape[0]): # add the rest of the data frame
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        k+=1
        t.style = 'TableGrid'
        

    #Save document in reports folder
    document.save(os.path.join(path, 'zoneAnomaly_report.docx'))
    convert(os.path.join(path, 'zoneAnomaly_report.docx'), os.path.join(path, 'report.pdf'))

    #Remove all used files
    os.remove(os.path.join(path,'zone_heat_season.png'))
    os.remove(os.path.join(path,'zone_cool_season.png'))
    os.remove(os.path.join(path,'zone_anomaly_summary.xlsx'))
    os.remove(os.path.join(path,'zoneAnomaly_report.docx'))

    print('Report successfully generated!')

    return

#------------------------------------------------------GENERATE END-USE DISAGGREGATION REPORT -------------------------------------------------
def endUseDisaggregation(path):
    print('Generating report for end-uses disaggregation function...')

    #Extract KPIs excel sheet
    kpis_htg = pd.read_excel(os.path.join(path,'endUseDisagg_summary.xlsx'),sheet_name='Heating')
    kpis_clg = pd.read_excel(os.path.join(path,'endUseDisagg_summary.xlsx'),sheet_name='Cooling')
    kpis_ele = pd.read_excel(os.path.join(path,'endUseDisagg_summary.xlsx'),sheet_name='Electricity')
    kpis_htg.drop(kpis_htg.columns[0],axis=1,inplace=True)
    kpis_clg.drop(kpis_clg.columns[0],axis=1,inplace=True)
    kpis_ele.drop(kpis_ele.columns[0],axis=1,inplace=True)

    #Generate Word document
    document = Document()

    #Report title
    document.add_heading('End-use Disaggregation - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The end-use disaggregation function calculates ') 
    p.add_run('annual energy use intensities of major end-uses. ').bold = True
    p.add_run('The major end-uses for electricity are lighting and plug-loads, distribution (i.e., pumps and fans), and chillers. \
The major end-uses for heating energy use are perimeter heating, the AHUs’ heating coils, and other appliances (i.e., domestic \
hot water). The major end-use for cooling energy are the AHUs’ cooling coils. This function can help assess \
the distribution of energy consumption within a building and can be used to identify abnormal energy use patterns.')

    #Visualization heading and description
    document.add_heading('Visualizations', level=1)
    p = document.add_paragraph('The visuals plot energy use intensity profiles which depict the weekly distribution \
of the major end-uses for electricity, cooling, and heating separately.')

    #Add visualizations
    document.add_picture(os.path.join(path,'endUseDisaggregation.png'), width=Inches(5.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('This section presents the generated KPIs - ')
    p.add_run('energy use intensities for major end-uses in heating energy, cooling energy, and electricity use').bold = True
    p.add_run(". The calculated energy use intensities for the AHU heating and cooling coils are disaggregated by each AHU. For \
example, if three (3) AHUs were analyzed, the energy use intensites for the AHU heating coils is provided as [NUM NUM NUM], whereby \
'NUM' represents the calculated energy use intensities for each consecutive AHU.")

    document.add_page_break()

    #Add KPIs
    table_labels = ['End-use',['Annual space heating load intensity (kWh/m2)','Annual space cooling load intensity (kWh/m2)','Annual electricity use intensity (kWh/m2)']]
    table_headings = ['Heating energy','Cooling energy','Electricity']
    k = 0

    for df in [kpis_htg,kpis_clg,kpis_ele]:
        df = df.round(1)
        p = document.add_heading(table_headings[k],level=2)
        t = document.add_table(df.shape[0]+1, df.shape[1])
        
        for j in range(df.shape[-1]): # add the header rows.
            if j == 0:
                t.cell(0,j).text = table_labels[j]
            else:
                t.cell(0,j).text = table_labels[j][k]
        
        for i in range(df.shape[0]): # add the rest of the data frame
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        k+=1
        t.style = 'Colorful List'

    #Save document in reports folder
    document.save(os.path.join(path,'endUseDisaggregation_report.docx'))
    convert(os.path.join(path,'endUseDisaggregation_report.docx'), os.path.join(path,'report.pdf'))
    print('Report successfully generated!')

    # remove used files
    os.remove(os.path.join(path,'endUseDisaggregation.png'))
    os.remove(os.path.join(path,'endUseDisaggregation_report.docx'))
    os.remove(os.path.join(path,'endUseDisagg_summary.xlsx'))

    return

#------------------------------------------------------GENERATE OCCUPANCY REPORT-------------------------------------------------
def occupancy(path,is_wifi):

    if is_wifi:

        print('Generating report for occupancy function with Wi-Fi device count data...')

        #Extract KPIs excel sheet
        kpis = pd.read_excel(os.path.join(path,'arrive_depart_maxOcc.xlsx'),sheet_name='KPIs',keep_default_na=False)

        #Generate Word document
        document = Document()

        #Report title
        document.add_heading('Occupancy (Wi-Fi device count) - Analysis Report', 0)

        #Introductory paragraph
        p = document.add_paragraph('The occupancy function determines typical ')
        p.add_run('earliest arrival times, latest departure times,').bold = True
        p.add_run(' and ')
        p.add_run('highest recorded occupancy').bold = True
        p.add_run(' for weekdays and weekends separately. This function can be used to inform ventilation schedules which can minimize \
excessive ventilation during unoccupied hours, or even serve as a bssis for an occupant-driven demand controlled ventilation scheme. \
Visuals plot building-level and floor-level* occupant count profiles.')
        p = document.add_paragraph("*Floor-level occupant count profiles are shown only if multiple files were inputted and read.")

        #Visualization heading and description - Building-level occupant count
        document.add_heading('Visuals - Building-level occupant count', level=1)
        p = document.add_paragraph('This set of visuals plots hourly building-level occupant count profiles for weekdays (to the left) and \
weekends (to the right) separately. The profiles are shown as the 25th, 50th (median), and 75th percentile. The 25th percentile represents \
the lower range of typical occupancy and the 75th represents the higher range. The second set of visualizations are \
occupant count profiles taken at the 75th percentile and broken down by floor.')

        #Add visualizations
        document.add_picture(os.path.join(path,'percentile_occ.png'), width=Inches(5.75))

        if os.path.isfile(os.path.join(path, 'floor_level_occ.png')):

            #Visualization heading and description - Floor-level occupant count
            document.add_heading('Visuals - Floor-level occupant count', level=1)
            p = document.add_paragraph('This set of visuals plots hourly floor-level occupant count profiles for weekdays (to the left) and \
weekends (to the right) separately. The profiles are broken down by floor and are shown as the 75th, which is the higher range of typical occupancy.')

            #Add visualizations and remove after used
            document.add_picture(os.path.join(path, 'floor_level_occ.png'), width=Inches(5.75))
            os.remove(os.path.join(path,'floor_level_occ.png'))

        #KPIs heading and description
        document.add_heading('Key performance Indicators', level=1)
        p = document.add_paragraph('This section presents the generated KPIs - ')
        p.add_run('typical earliest arrival times, latest departure times, and highest occupant count').bold = True
        p.add_run(' which are broken down by floor and determined sparately for weekdays and weekends. The typical earliest arrival time \
is the hour when the occupant count exceeds '+r'10%'+' of the maximum recorded count per floor. Similiarily, the typical latest \
departure time is the hours when occupant count dips below '+r'10%'+' of the maximum recorded count per floor. The highest occupant count \
is taken as the highest recorded occupant count. Note all these calculations are determined at the 75th percentile, meaning the higher range \
of typical occupancy.')

        #Add tables with summary tables
        table_labels = ['Floor',
                        'WEEKDAY earliest arrival time',
                        'WEEKDAY latest departure time',
                        'WEEKDAY highest occupant count',
                        'WEEKEND earliest arrival time',
                        'WEEKEND latest departure time',
                        'WEEKEND highest occupant count']

        p = document.add_heading('Earliest arrival times, latest departure times, and highest occupant count',level=2)
        t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
        
        for j in range(kpis.shape[-1]): # add the header rows.
            t.cell(0,j).text = table_labels[j]
        
        for i in range(kpis.shape[0]): # add the rest of the data frame
            for j in range(kpis.shape[-1]):
                t.cell(i+1,j).text = str(kpis.values[i,j])
        
        t.style = 'Colorful List'

        #Save document in reports folder
        document.save(os.path.join(path,'occupancy_wifi_report.docx'))
        convert(os.path.join(path,'occupancy_wifi_report.docx'), os.path.join(path,'report.pdf'))
        print('Report successfully generated!')

        #Remove all used files
        os.remove(os.path.join(path,'occupancy_wifi_report.docx'))
        os.remove(os.path.join(path,'percentile_occ.png'))
        os.remove(os.path.join(path,'arrive_depart_maxOcc.xlsx'))
    
    else:

        print('Generating report for occupancy function with motion-detection data...')

        # get working directory
        cwd = os.getcwd()

        #Extract KPIs excel sheet
        kpis = pd.read_excel(os.path.join(path,'motion_detection_kpis.xlsx'),sheet_name='KPIs',keep_default_na=False)
        kpis.drop(kpis.columns[0],axis=1,inplace=True)

        #Generate Word document
        document = Document()

        #Report title
        document.add_heading('Occupancy (Motion-detection) - Analysis Report', 0)

        #Introductory paragraph
        p = document.add_paragraph('The occupancy function determines the typical ')
        p.add_run('earliest arrival time, latest arrival time, latest departure time,').bold = True
        p.add_run(' and ')
        p.add_run('longest break duration').bold = True
        p.add_run(' for weekdays only.')
        p = document.add_paragraph("This function can be used to inform ventilation schedules which can minimize \
excessive ventilation during unoccupied hours, or even serve as a basis for an occupant-driven demand controlled ventilation scheme. \
The calculated earliest arrival time and latest departure time can be used to inform whole building-level schedules, including start/stop \
times, whereas the calculated latest arrival time and longest break duration can be used to inform zone-level schedules, including time-out \
durations for zone air flow. For example, if the calculated latest arrival time is 13:00 (1PM), you may implement a simple logic to each \
zone's controller which switches the zone to the 'unoccupied mode' when the motion detector of the zone has not been triggered until 1 pm \
on a given day. Similarly, if the calculated longest break duration is 3:00 (3 hours), then you may implement another simple logic to each \
zone's controller to switch the zone to the 'unoccupied mode' when the room's motion detector has not been triggered for longer than 3 hours \
on a given day. The below pseudo-code demonstrates these two implementations.")

        #KPIs heading and description
        document.add_heading('Key performance Indicators', level=1)
        p = document.add_paragraph('This section presents the generated KPIs - ')
        p.add_run('typical earliest arrival time, latest arrival time, latest departure time, and longest break duration').bold = True
        p.add_run(' for weekdays only.')

        #Add tables with summary tables
        table_labels = ['Earliest arrival time',
                        'Latest arrival time',
                        'Latest departure time',
                        'Longest break duration (hours)'
                        ]

        t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
        
        for j in range(kpis.shape[-1]): # add the header rows.
            t.cell(0,j).text = table_labels[j]
        
        for i in range(kpis.shape[0]): # add the rest of the data frame
            for j in range(kpis.shape[-1]):
                t.cell(i+1,j).text = str(kpis.values[i,j])
        
        t.style = 'Colorful List'

        #Save document in reports folder
        document.save(os.path.join(path,'occupancy_motion_report.docx'))
        convert(os.path.join(path,'occupancy_motion_report.docx'), os.path.join(path,'report.pdf'))

        #Remove all used files
        os.remove(os.path.join(path,'occupancy_motion_report.docx'))
        os.remove(os.path.join(path,'motion_detection_kpis.xlsx'))
    
    print('Report successfully generated!')

    return

#------------------------------------------------------GENERATE COMPLAINTS REPORT-------------------------------------------------
def complaintAnalytics(path):
    print('Generating report for occupant complaints analytics function...')

    #Extract KPIs excel sheet
    kpis = pd.read_excel(os.path.join(path,'complaints_freq.xlsx'),sheet_name='Daily frequency of complaints')
    kpis.drop(kpis.columns[0],axis=1,inplace=True)
    kpis.iloc[:,1:] = kpis.iloc[:,1:].astype(float).round(3)

    #Generate Word document
    document = Document()

    #Report title
    document.add_heading('Complaint Analytics - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The occupant complaints function determines the ')
    p.add_run('daily frequency of hot/cold-related occupant complaints. ').bold = True
    p.add_run('This function can be used to assess the effects of temperature setpoint or \
schedule adjustments on building occupants. The visuals depict the distribution of complaints by month and by the period \
of the day, the relationship between the complaints and indoor and outdoor air temperature, and the predicted proportion \
of the complaints based on time of day, day of the week, and outdoor air temperature. More information is available at the \
respective sections.')

    #First visualization heading and description
    document.add_heading('Visuals - Occupant complaints breakdown', level=1)
    p = document.add_paragraph('This set of visuals categorizes the complaints by the type of complaint \
(hot or cold related), and counts the number of complaints by the month and period of the day the complaint was \
registered.')

    #Add first visualizations
    document.add_picture(os.path.join(path,'complaints_breakdown.png'), width=Inches(5.75))

    #Second visualization heading and description
    document.add_heading('Visuals - Indoor and outdoor air temperature', level=1)
    p = document.add_paragraph('This visual plots the relationship between a hot/cold complaint \
and the indoor and outdoor air temperature at the time the complaint was registered.')

    #Add second visualizations
    document.add_picture(os.path.join(path,'complaint_scatter.png'), width=Inches(5.75))

    #Third visualization heading and description
    document.add_heading('Visuals - Predicted proportion of complaints', level=1)
    p = document.add_paragraph("This visual predicts the proportion of complaints that would be made \
with respect to certain conditions. The boxes which branch to the left represent the predicted proportion of complaints when \
the condition in the preceding box is satisfied. For example, if a box with the condition, 'Hour of the day <= 10' \
has a left branch with "+r"40%"+" and a right branch with "+r"60%"+", it is predicted that "+r"40%"+" of all complaints \
made will occur before 10 am. The condition is displayed at the top of the box, and the predicted \
proportion is displayed at the center of each box between the 0.0s.")

    #Add third visualizations
    document.add_picture(os.path.join(path,'decision_tree.png'), width=Inches(5.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('This section presents the generated KPIs - ')
    p.add_run('daily frequency of hot and cold complaints in the heating and cooling season.').bold = True
    p.add_run(' The values represent the number of complaints made per day. Higher frequencies indicate a higher \
rate of occurence of a type of complaint for the particular season.')

    table_labels = ['Type of complaint',
                    'Daily frequency for heating season (complaints per day)',
                    'Daily frequency for cooling season (complaints per day)']

    p = document.add_heading('Daily frequencies of hot/cold occupant complaints',level=2)
    t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
    
    for j in range(kpis.shape[-1]): # add the header rows.
        t.cell(0,j).text = table_labels[j]
    
    for i in range(kpis.shape[0]): # add the rest of the data frame
        for j in range(kpis.shape[-1]):
            t.cell(i+1,j).text = str(kpis.values[i,j])
    
    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(os.path.join(path,'complaint_report.docx'))
    convert(os.path.join(path,'complaint_report.docx'), os.path.join(path,'report.pdf'))

    #remove all used files
    os.remove(os.path.join(path,'complaints_freq.xlsx'))
    os.remove(os.path.join(path,'complaints_breakdown.png'))
    os.remove(os.path.join(path,'complaint_scatter.png'))
    os.remove(os.path.join(path,'decision_tree.png'))
    os.remove(os.path.join(path,'complaint_report.docx'))

    print('Report successfully generated!')

    return

#METADATA

def metadata(path):
    print('Generating report for metadata function...')

    #Extract KPIs excel sheet
    ahu_meta = pd.read_excel(os.path.join(path,'metadata_summary.xlsx'),sheet_name='ahu')
    zone_meta = pd.read_excel(os.path.join(path,'metadata_summary.xlsx'),sheet_name='zone')
    ahu_meta.drop(ahu_meta.columns[0],axis=1,inplace=True)
    zone_meta.drop(ahu_meta.columns[0],axis=1,inplace=True)

    #Generate Word document in landscape orientation
    document = Document()
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE

    #Report title
    document.add_heading('Metadata - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The metadata inference function automatically ')
    p.add_run('identifies and associates BAS metadata labels into corresponding data point types and AHU or zone. ').bold = True
    p.add_run('This function is unique as it is intended as a pre-processing step rather than the other functions within the \
toolkit which are intended to help the user identify energy deficiencies and anomalies and oppurtunities for energy-savings. \
This function does not output any key performance indicators (KPI) or visuals like the other functions in the toolkit. Rather, \
it outputs a set of metadata labels which are organized by its data type (i.e., temperature sensor, \
actuator position, valve position) and by system hierarchy (AHU or zone). Note that this function works by identifying labels based on \
a predefined dictionary of common abbreviations for certain data points (i.e., "TSA" or "SAT" for supply air temperature sensor) and \
associates them based on the alphanumeric similiarity of certain labels to other labels. Hence, this function may not work as effectively \
with BAS labels with unique ontologies and inconsistent labelling schema.')

    #Output metadata labels for AHUs
    document.add_heading('Metadata labels', level=1)
    table_labels = ['AHU',
                    'tSa label',
                    'tSa ID',
                    'tRa label',
                    'tRa ID',
                    'tOa label',
                    'tOa ID',
                    'pSa label',
                    'pSa ID',
                    'sOa label',
                    'sOa ID',
                    'sHc label',
                    'sHc ID',
                    'sCc label',
                    'sCc ID',
                    'sFan label',
                    'sFan ID']

    p = document.add_heading('Metadata label for AHUs',level=2)
    t = document.add_table(ahu_meta.shape[0]+1, ahu_meta.shape[1])
    
    for j in range(ahu_meta.shape[-1]): # add the header rows.
        t.cell(0,j).text = table_labels[j]
    
    for i in range(ahu_meta.shape[0]): # add the rest of the data frame
        for j in range(ahu_meta.shape[-1]):
            t.cell(i+1,j).text = str(ahu_meta.values[i,j])
    
    t.style = 'Colorful List'

    #Output metadata labels for zones
    table_labels = ['Zone',
                    'tIn label',
                    'tIn ID',
                    'qFlo label',
                    'qFlo ID',
                    'qFloSp label',
                    'qFloSp ID',
                    'sRad label',
                    'sRad ID']

    p = document.add_heading('Metadata label for zones',level=2)
    t = document.add_table(zone_meta.shape[0]+1, zone_meta.shape[1])
    
    for j in range(zone_meta.shape[-1]): # add the header rows.
        t.cell(0,j).text = table_labels[j]
    
    for i in range(zone_meta.shape[0]): # add the rest of the data frame
        for j in range(zone_meta.shape[-1]):
            t.cell(i+1,j).text = str(zone_meta.values[i,j])
    
    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(os.path.join(path,'meta_report.docx'))
    convert(os.path.join(path,'meta_report.docx'), os.path.join(path,'report.pdf'))

    #remove all used files
    os.remove(os.path.join(path,'metadata_summary.xlsx'))
    os.remove(os.path.join(path,'meta_report.docx'))

    print('Report successfully generated!')

    return