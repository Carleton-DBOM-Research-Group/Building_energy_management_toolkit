#import required modules
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

#-----------------------------------------------GENERATE BASELINE ENERGY REPORT----------------------------------------------------
def energyBaseline():
    print('Generating report for Baseline Energy function...')

    path = os.getcwd() #Get current directory
    output_path = path + r'\reports\2-energyBaseline' #Specify the output path for report
    input_path = path + r'\outputs\2-energyBaseline'

    document = Document()

    document.add_heading('Baseline Energy - Analysis Report', 0) #Report title

    #Introductory paragraph
    p = document.add_paragraph('The baseline energy function ')
    p.add_run('compares energy use during operating hours (workhours) and outside operating hours (after-hours) for heating, cooling, and electricity. ').bold=True
    p.add_run("This function is intended to help the user assess the effectiveness of schedules and their ability to reduce \
energy use outside of the building's operating hours. Plots are generated which compare the rate of energy use during and outside operating \
hours with respect to outdoor air temperature, and predict energy consumption at representative outdoor air temperatures - these are \
done separately for heating, cooling, and electrcity. The generated key performance indicators (KPI) quantify schedule \
effectiveness and afterhours energy use. More information is found in the respective sections.")

    #Visualization heading and description
    document.add_heading('Visuals', level=1)
    p = document.add_paragraph('The first set of visuals (to the left) compare the rate of energy use during operating hours (workhours) and \
outside operating hours (after-hours) as a function of outdoor air temperatures - this is done separately for heating, \
cooling, and electricity. Current schedules may be ineffective at reducing after-hours energy use if the after-hours energy use \
line is similiar or identical to the workhours energy use line. If the slope of the after-hours energy use line is shallower \
than the workhour eneegry use line, the current schedules are effectively reducing energy use outside operating hours.')
    p = document.add_paragraph('The second set of visuals (to the right) illustrates the sensitivity of energy use with respect to outdoor \
air temperatures - this is done separately for heating, cooling, and electricity. If the lines are spaced considerably apart, \
the energy use is particularily sensitive to outdoor air temperature.')

    #Add visualizations
    document.add_picture(input_path + r'\energyBase_heating.png', width=Inches(5.75))
    document.add_picture(input_path + r'\energyBase_cooling.png', width=Inches(5.75))
    document.add_picture(input_path + r'\energyBase_electricity.png', width=Inches(5.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('The key performance indicators (KPIs) are ')
    p.add_run('Schedule effectiveness').bold = True
    p.add_run(' and ')
    p.add_run('Afterhours energy use ratio').bold = True
    p.add_run('. Schedule effectiveness quantifies the difference between the slope of the workhours energy use line and the \
afterhours energy use line. Values approaching zero (0%) indicate similiar or identical inclined slopes, positive (+) values indicate \
a steeper workhours energy use slope than afterhours, and negative (-) values indicate a steeper afterhours energy use slope. Therefore, \
a greater positive value is desirable since it indicates an effective reduction in energy use rates during afterhours.\n\n\
The Afterhours energy use ratio is the ratio of energy use during afterhours over the total energy use. Higher value indicate a larger \
portion of total energy consumption used during after-hours. Therefore, a lower value is desirable.')

    #Extract data from excel sheet and add table to document
    kpis = pd.read_excel(input_path + r'\energyBase_summary.xlsx',sheet_name='KPIs')
    kpis.drop(kpis.columns[0],axis=1,inplace=True)
    kpis.iloc[:,1:] = (kpis.iloc[:,1:] * 100).astype(int).astype(str) + '%'

    p = document.add_heading('Schedule Effectiveness and Afterhours energy use ratio',level=2)
    t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
    
    for j in range(kpis.shape[-1]): # add the header rows.
        t.cell(0,j).text = kpis.columns[j]
    
    for i in range(kpis.shape[0]): # add the rest of the data frame
        for j in range(kpis.shape[-1]):
            t.cell(i+1,j).text = str(kpis.values[i,j])
    
    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(output_path + r'\energyBaseline_report.docx')
    convert(output_path + r'\energyBaseline_report.docx', output_path + r'\energyBaseline_report.pdf')
    print('Report successfully generated!')

    return

#-----------------------------------------------GENERATE AHU ANOMALY REPORT----------------------------------------------------
def ahuAnomaly():
    print('Generating report for AHU anomaly detection function...')

    path = os.getcwd() #Get current directory
    output_path = path + r'\reports\3-ahuAnomaly' #Specify the output path for report
    input_path = path + r'\outputs\3-ahuAnomaly'

    #Extract KPIs excel sheet
    faults_df = pd.read_excel(input_path + r'\ahu_anomaly_summary.xlsx',sheet_name='faults')
    faults_df.drop(faults_df.columns[0],axis=1,inplace=True)
    faults_df.iloc[:,1] = (faults_df.iloc[:,1]).astype(int).astype(str) + '%'

    document = Document()

    document.add_heading('AHU Anomaly - Analysis Report', 0) #Report title

    #Introductory paragraph
    p = document.add_paragraph('The AHU anomaly detection function ')
    p.add_run('detects hard and soft faults related to air handling units (AHUs). ').bold=True
    p.add_run("This function helps identify potential causes for anomalous AHU operations which may cause \
energy use inefficiencies. The visuals are intended to aid understanding of the detected faults. These depict \
supply air temperature, and the coolest/warmest/average return air temperatures as a function of outdoor air temperature, and \
damper and valve actuator positions as a function of outdoor air temperature. Additionally, a number of diagrams are generated which \
depict damper and valve actuator positions and temperature readings at characteristic AHU operational periods. More information is \
available at the respective sections.")

    #Visualization heading and description - Part 1
    document.add_heading('Visuals - Split-range controller', level=1)
    p = document.add_paragraph('A set of two charts are generated for each AHU dataset inputted. The first (top) plots supply air \
temperature, and the coolest/warmest/average return air temperatures as a function of outdoor air temperature. For \
reference, the "ideal" supply air temperature is depicted. The second (bottom) chart is a Split-range controller diagram, which \
plots the outdoor air damper position (OA), heating coil valve position (HC), cooling coil valve position (CC) and average fraction \
of active perimeter heaters (RAD) with respect to outdoor air temperature. The four underlaying color zones represent the four \
distinct operating mode: Heating (red zone), economizer (yellow zone), economizer with cooling (grey zone), and cooling (blue zone). \
As an example, the below Split-range controller diagram is representative of normal, healthy AHU operations. Some ')
    p=document.add_paragraph('key characteristics of normal, healthy AHU operation').bold=True
    p=document.add_paragraph(' include:')

    document.add_paragraph('Heating coil active ONLY in heating mode', style='List Bullet')
    document.add_paragraph('Cooling coil active ONLY in economizer with cooling and cooling mode', style='List Bullet')
    document.add_paragraph('Heating and cooling coils should not operate simutaneously', style='List Bullet')
    document.add_paragraph('Perimeter heating should be minimal in economizer mode. This can be achieved by increasing the supply air \
temperature setpoint in the heating season, while monitoring any occurence of overheating.', style='List Bullet')

    document.add_picture(output_path + r'\SRC_example.png', width=Inches(5.75))

    p = document.add_paragraph('Suboptimal supply air temperatures can result excessive energy consumption from excessive perimeter heating use, \
economizing, or fan use. To guide supply air temperature setpoint adjustments, the typical "ideal" supply air temperature range is provided as a \
reference. If the supply air temperature is LOWER than this range in the heating season, excessive use of perimeter heating and economizing may \
result; only a few overheating rooms may trigger this behaviour. If this is the case, consider increasing the maximum terminal airflow setpoints \
in these overheating rooms.\n\nIf the supply air temperature is HIGHER than this range in the cooling season, excessive fan power to deliver \
required necessary cooling may result; a few overcooling rooms may trigger this behaviour. If this is the case, consider decreasing the minimum \
terminal airflow setpoints in these overcooling rooms within reason. In both of these cases, ensure that the airflow and temperature sensors work \
as intended in these rooms.')

    #For each AHU, output the first set of visuals
    for i in faults_df.index:

        document.add_heading('AHU: ' + str(faults_df.iloc[i][0]), level=2)
        document.add_picture(input_path + r'\f2a_ahu_'+str(i+1)+'.png', width=Inches(5))
    
    #Visualization heading and description - Part 2
    document.add_heading('Visuals - AHU operating periods', level=1)
    p = document.add_paragraph("A set of four to six visuals per AHU are generated which depict characteristic operating \
periods of the AHU and the average damper and valve positions and temperatures at those periods. The fraction of time of \
operation is the percentage of the total time of the AHU's operation which exhibit the displayed damper/valve positions \
and temperatures. Some ")
    p=document.add_paragraph('key characteristics of normal, healthy AHU operation').bold=True
    p=document.add_paragraph(' include:')

    document.add_paragraph('Heating coil active ONLY when the outdoor air temperature is below the upper setpoint for heating mode. (In the example \
split-range controller diagram provided, this is -5C.)', style='List Bullet')
    document.add_paragraph('Cooling coil active ONLY when the outdoor air temperature is above the upper setpoint for economizer mode. (In the \
example split-range controller diagram provided, this is 10C.)', style='List Bullet')
    document.add_paragraph('Heating and cooling coils should not operate simutaneously', style='List Bullet')
    document.add_paragraph('Perimeter heating should be minimal in economizer mode (i.e. when both the heating and cooling \
coils are inactive). This can be achieved by increasing the supply air temperature setpoint in the heating season, while \
monitoring any occurence of overheating.', style='List Bullet')
    
    #For each AHU, output the second set of visuals
    for i in faults_df.index:

        document.add_heading('AHU: ' + str(faults_df.iloc[i][0]), level=2)

        for j in range(1,7):
            try:
                document.add_picture(input_path + r'\f2b_ahu_'+str(i+1)+'_C_'+str(j)+'.png', width=Inches(5))
            except FileNotFoundError:
                break

    document.add_page_break()

    #Extract data from excel sheet and add table to document

    p = document.add_heading('Key performance indicators - AHU faults',level=1)
    p = document.add_paragraph('The following table lists the hard and soft faults identified by the function. The AHU health \
index is also provided for each AHU which is '+ r'100%'+' if no faults are detected and '+ r'0%'+' if all six faults are detected. \
The following faults may be detected:')

    document.add_paragraph('Cooling coil stuck: This fault is generated if no or minimal use of the cooling coil was observed in the data. \
This may be symptomatic of a faulty sensor, a stuck valve, or a conflict in the operational logic.', style='List Bullet')
    document.add_paragraph('Heating coil stuck: This fault is generated if no or minimal use of the cooling coil was observed in the data. \
This may be symptomatic of a faulty sensor, a stuck valve, or a conflict in the operational logic.', style='List Bullet')
    document.add_paragraph("Check economizer logic: This fault is generated if excessive use of the perimeter heatings was observed in the \
AHU's economizer mode.", style='List Bullet')
    document.add_paragraph('Low/High outdoor air: This fault is generated if an inadequate or excessive amount of outdoor air was observed. \
This may be symptomatic of a stuck outdoor air damper or faulty damper sensor/actuator.', style='List Bullet')
    document.add_paragraph('Check mode of operation logic: This fault is generated if the weekly operational time (i.e. when the AHU fans are \
operational) exceeds 100 hours a week. This considered excessive operations. It is suggested to check the operational logic for any conflicts \
which may result in unintended operation of the AHUs', style='List Bullet')
    document.add_paragraph("Check supply air temperature reset logic: This fault is generated if there is excessive use of perimeter heating \
devices during the AHUs' economizer mode. This may be a result of select rooms overheating in the heatinmg season. If this is the case, consider \
increasing the maximum terminal airflow setpoints in these overheating rooms.", style='List Bullet')

    t = document.add_table(faults_df.shape[0]+1, faults_df.shape[1])
    
    for j in range(faults_df.shape[-1]): # add the header rows.
        t.cell(0,j).text = faults_df.columns[j]
    
    for i in range(faults_df.shape[0]): # add the rest of the data frame
        for j in range(faults_df.shape[-1]):
            t.cell(i+1,j).text = str(faults_df.values[i,j])

    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(output_path + r'\ahuAnomaly_report.docx')
    convert(output_path + r'\ahuAnomaly_report.docx', output_path + r'\ahuAnomaly_report.pdf')
    print('Report successfully generated!')

    return

#-----------------------------------------------GENERATE ZONE ANOMALY REPORT----------------------------------------------------
def zoneAnomaly():
    print('Generating report for zone anomaly detection function...')

    path = os.getcwd() #Get current directory
    output_path = path + r'\reports\4-zoneAnomaly' #Specify the output path for report
    input_path = path + r'\outputs\4-zoneAnomaly' #Specify the input path for report

    #Extract KPIs excel sheet
    kpis_heating = pd.read_excel(input_path + r'\zone_anomaly_summary.xlsx',sheet_name='Heating season')
    kpis_heating.drop(kpis_heating.columns[0],axis=1,inplace=True)

    kpis_cooling = pd.read_excel(input_path + r'\zone_anomaly_summary.xlsx',sheet_name='Cooling season')
    kpis_cooling.drop(kpis_cooling.columns[0],axis=1,inplace=True)

    #Generate Word document
    document = Document()

    #Report title
    document.add_heading('Zone Anomaly - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The zone anomaly function ')
    p.add_run('detects anomalous zones based on indoor air temperature and airflow control errors. ').bold = True
    p.add_run('This function can help identify potential faults in variable air volume (VAV) units which \
may result in anomalous airflow and temperature conditions in thermal zones. Visuals are also generated which \
plot the average indoor air temperature and airflow control error or groups of zones; this is done separately \
for the heating and cooling season.')

    #Visualization heading and description
    document.add_heading('Visuals', level=1)
    p = document.add_paragraph('The generated visuals plot the average indoor air temperature and average \
airflow control error or groups of zones. Each group of zones (C1, C2, C3, etc.) represents a number of zones \
which is presented at the top-left of the visual. The airflow control error is the difference between the \
actual and setpoint airflow rate with respect to the airflow rate setpoint. A positive (+) percentage indicates a \
higher actual flowrate than the setpoint and a negative (-) percentage indicates a lower actual flowrate than the \
setpoint. This visual is generated twice, once for the heating season (December through February) and again \
for the cooling season (May through August). The following are possible symptoms of anomalous operations:')

    document.add_paragraph('High flow: Zones exhibit abnormally high airflow. Ensure the VAV termianl damper and \
airflow sensors are operating as intended. High airflow to zones may result in excessive use of perimeter heaters.', style='List Bullet')
    document.add_paragraph('Low flow: Zones exhibit abnormally low airflow. Ensure the VAV termianl damper and \
airflow sensors are operating as intended. Low airflow to zones may result in inadequate indoor air quality for occuapnts.', style='List Bullet')
    document.add_paragraph("Cold: Zones exhibit abnormally low indoor air temperatures. If zones exhibit acceptable \
airflow control errors, ensure perimeter heating devices and/or reheat coils are operating as intended. Consider \
decreasing the minimum airflow setpoint.", style='List Bullet')
    document.add_paragraph("Hot: Zones exhibit abnormally high indoor air temperatures. If zones exhibit acceptable \
airflow control errors, ensure perimeter heating devices and/or reheat coils are operating as intended. Consider \
increasing the maximum airflow setpoint.", style='List Bullet')

    #Add visualizations
    document.add_picture(input_path + r'\zone_heat_season.png', width=Inches(4.75))
    document.add_picture(input_path + r'\zone_cool_season.png', width=Inches(4.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('This section presents the generated KPIs - ')
    p.add_run('the zone health index').bold = True
    p.add_run('. The zone health index is the ') 
    p.add_run('ratio of zones with acceptable indoor air temperature and airflow control error over the \
total number of zones ').bold = True
    p.add_run('- this is calculated separately for the heating season (December through February) and the \
cooling season (May through August). An acceptable indoor air temperature is considered to be between 20 and \
25 degrees, and an acceptable airflow control error is +/- 20%. A greater percentage is desirable since this \
indicates little to no detected anomalous zones.')

    heating_health_index = (round(kpis_heating.iloc[0]['Health Index'],3))*100
    cooling_health_index = (round(kpis_cooling.iloc[0]['Health Index'],3))*100

    p = document.add_paragraph('Zone health index for heating season: ', style='List Bullet')
    p.add_run(str(heating_health_index) + '%').bold = True
    p = document.add_paragraph('Zone health index for cooling season: ', style='List Bullet')
    p.add_run(str(cooling_health_index) + '%').bold = True

    #Add tables with summary tables
    p = document.add_paragraph('The following tables lists the number of zones in each group and the average \
indoor air temperature and average airflow control error for each group. For the heating season, the average fraction \
of active (ON-state) perimeter heaters within a zone is also provided for each cluster.')
    
    document.add_page_break()

    kpis_heating.drop(kpis_heating.columns[-1],axis=1,inplace=True) #Drop the health index from the heating season table
    kpis_heating.iloc[:,0] = (kpis_heating.iloc[:,0]).astype(int)
    kpis_heating.iloc[:,2:] = (kpis_heating.iloc[:,2:]).round(1).astype(str) + '%'

    kpis_cooling.drop(kpis_cooling.columns[-1],axis=1,inplace=True) #Drop the health index from the cooling season table
    kpis_cooling.iloc[:,0] = (kpis_cooling.iloc[:,0]).astype(int)
    kpis_cooling.iloc[:,2] = (kpis_cooling.iloc[:,2]).round(1).astype(str) + '%'

    table_labels = ['Number of zones','Average indoor air temperature (C)','Average airflow control error','Average fraction of active perimeter heaters']
    table_headings = ['Heating season KPIs','Cooling season KPIs']
    k = 0

    for df in [kpis_heating,kpis_cooling]:
        df = df.round(1)
        p = document.add_heading(table_headings[k],level=2)
        t = document.add_table(df.shape[0]+1, df.shape[1])
        
        for j in range(df.shape[-1]): # add the header rows.
            t.cell(0,j).text = table_labels[j]
        
        for i in range(df.shape[0]): # add the rest of the data frame
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        k+=1
        t.style = 'Colorful List'
        

    #Save document in reports folder
    document.save(output_path + r'\zoneAnomaly_report.docx')
    convert(output_path + r'\zoneAnomaly_report.docx', output_path + r'\zoneAnomaly_report.pdf')
    print('Report successfully generated!')

    return

#------------------------------------------------------GENERATE END-USE DISAGGREGATION REPORT -------------------------------------------------
def endUseDisaggregation():
    print('Generating report for end-uses disaggregation function...')

    path = os.getcwd() #Get current directory
    output_path = path + r'\reports\5-endUseDisaggregation' #Specify the output path for report
    input_path = path + r'\outputs\5-endUseDisaggregation' #Specify the input path for report

    #Extract KPIs excel sheet
    kpis_htg = pd.read_excel(input_path + r'\endUseDisagg_summary.xlsx',sheet_name='Heating')
    kpis_clg = pd.read_excel(input_path + r'\endUseDisagg_summary.xlsx',sheet_name='Cooling')
    kpis_ele = pd.read_excel(input_path + r'\endUseDisagg_summary.xlsx',sheet_name='Electricity')
    kpis_htg.drop(kpis_htg.columns[0],axis=1,inplace=True)
    kpis_clg.drop(kpis_clg.columns[0],axis=1,inplace=True)
    kpis_ele.drop(kpis_ele.columns[0],axis=1,inplace=True)

    #Generate Word document
    document = Document()

    #Report title
    document.add_heading('End-use Disaggregation - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The end-use disaggregation function calculates ') 
    p.add_run('annual energy use intensities of major end-uses. ').bold = True
    p.add_run('The major end-uses for electricity are lighting and plug-loads, distribution (i.e., pumps and fans), and chillers. \
The major end-uses for heating energy use are perimeter heating, the AHUs’ heating coils, and other appliances (i.e., domestic \
hot water). The major end-use for cooling energy are the AHUs’ cooling coils. This function can help assess \
the distribution of energy consumption within a building and can be used to identify abnormal energy use patterns.')

    #Visualization heading and description
    document.add_heading('Visualizations', level=1)
    p = document.add_paragraph('The visuals plot energy use intensity profiles which depict the weekly distribution \
of the major end-uses for electricity, cooling, and heating separately.')

    #Add visualizations
    document.add_picture(input_path + r'\endUseDisaggregation.png', width=Inches(5.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('This section presents the generated KPIs - ')
    p.add_run('energy use intensities for major end-uses in heating energy, cooling energy, and electricity use').bold = True
    p.add_run(". The calculated energy use intensities for the AHU heating and cooling coils are disaggregated by each AHU. For \
example, if three (3) AHUs were analyzed, the energy use intensites for the AHU heating coils is provided as [NUM NUM NUM], whereby \
'NUM' represents the calculated energy use intensities for each consecutive AHU.")

    document.add_page_break()

    #Add KPIs
    table_labels = ['End-use',['Annual space heating load intensity (kWh/m2)','Annual space cooling load intensity (kWh/m2)','Annual electricity use intensity (kWh/m2)']]
    table_headings = ['Heating energy','Cooling energy','Electricity']
    k = 0

    for df in [kpis_htg,kpis_clg,kpis_ele]:
        df = df.round(1)
        p = document.add_heading(table_headings[k],level=2)
        t = document.add_table(df.shape[0]+1, df.shape[1])
        
        for j in range(df.shape[-1]): # add the header rows.
            if j == 0:
                t.cell(0,j).text = table_labels[j]
            else:
                t.cell(0,j).text = table_labels[j][k]
        
        for i in range(df.shape[0]): # add the rest of the data frame
            for j in range(df.shape[-1]):
                t.cell(i+1,j).text = str(df.values[i,j])
        k+=1
        t.style = 'Colorful List'

    #Save document in reports folder
    document.save(output_path + r'\endUseDisaggregation_report.docx')
    convert(output_path + r'\endUseDisaggregation_report.docx', output_path + r'\endUseDisaggregation_report.pdf')
    print('Report successfully generated!')

    return

#------------------------------------------------------GENERATE OCCUPANCY REPORT-------------------------------------------------
def occupancy(which_data):

    if which_data == 1:

        print('Generating report for occupancy function with Wi-Fi device count data...')

        path = os.getcwd() #Get current directory
        output_path = path + r'\reports\6-occupancy' #Specify the output path for report
        input_path = path + r'\outputs\6-occupancy' #Specify the input path for report

        #Extract KPIs excel sheet
        kpis = pd.read_excel(input_path + r'\arrive_depart_maxOcc.xlsx',sheet_name='KPIs',keep_default_na=False)

        #Generate Word document
        document = Document()

        #Report title
        document.add_heading('Occupancy (Wi-Fi device count) - Analysis Report', 0)

        #Introductory paragraph
        p = document.add_paragraph('The occupancy function determines typical ')
        p.add_run('earliest arrival times, latest departure times,').bold = True
        p.add_run(' and ')
        p.add_run('highest recorded occupancy').bold = True
        p.add_run(' for weekdays and weekends separately. This function can be used to inform ventilation schedules which can minimize \
    excessive ventilation during unoccupied hours, or even serve as a bssis for an occupant-driven demand controlled ventilation scheme. \
    Visuals plot building-level and floor-level occupant count profiles.')

        #Visualization heading and description - Building-level occupant count
        document.add_heading('Visuals - Building-level occupant count', level=1)
        p = document.add_paragraph('This set of visuals plots hourly building-level occupant count profiles for weekdays (to the left) and \
    weekends (to the right) separately. The profiles are shown as the 25th, 50th (median), and 75th percentile. The 25th percentile represents \
    the lower range of typical occupancy and the 75th represents the higher range. The second set of visualizations are \
    occupant count profiles taken at the 75th percentile and broken down by floor.')

        #Add visualizations
        document.add_picture(input_path + r'\percentile_occ.png', width=Inches(5.75))

        #Visualization heading and description - Floor-level occupant count
        document.add_heading('Visuals - Floor-level occupant count', level=1)
        p = document.add_paragraph('This set of visuals plots hourly floor-level occupant count profiles for weekdays (to the left) and \
    weekends (to the right) separately. The profiles are broken down by floor and are shown as the 75th, which is the higher range of typical occupancy.')

        #Add visualizations
        document.add_picture(input_path + r'\floor_level_occ.png', width=Inches(5.75))

        #KPIs heading and description
        document.add_heading('Key performance Indicators', level=1)
        p = document.add_paragraph('This section presents the generated KPIs - ')
        p.add_run('typical earliest arrival times, latest departure times, and highest occupant count').bold = True
        p.add_run(' which are broken down by floor and determined sparately for weekdays and weekends. The typical earliest arrival time \
    is the hour when the occupant count exceeds '+r'10%'+' of the maximum recorded count per floor. Similiarily, the typical latest \
    departure time is the hours when occupant count dips below '+r'10%'+' of the maximum recorded count per floor. The highest occupant count \
    is taken as the highest recorded occupant count. Note all these calculations are determined at the 75th percentile, meaning the higher range \
    of typical occupancy.')

        document.add_page_break()

        #Add tables with summary tables
        table_labels = ['Floor',
                        'WEEKDAY earliest arrival time',
                        'WEEKDAY latest departure time',
                        'WEEKDAY highest occupant count',
                        'WEEKEND earliest arrival time',
                        'WEEKEND latest departure time',
                        'WEEKEND highest occupant count']

        p = document.add_heading('Earliest arrival times, latest departure times, and highest occupant count',level=2)
        t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
        
        for j in range(kpis.shape[-1]): # add the header rows.
            t.cell(0,j).text = table_labels[j]
        
        for i in range(kpis.shape[0]): # add the rest of the data frame
            for j in range(kpis.shape[-1]):
                t.cell(i+1,j).text = str(kpis.values[i,j])
        
        t.style = 'Colorful List'

        #Save document in reports folder
        document.save(output_path + r'\occupancy_wifi_report.docx')
        convert(output_path + r'\occupancy_wifi_report.docx', output_path + r'\occupancy_wifi_report.pdf')
        print('Report successfully generated!')

        return
    
    else:

        print('Generating report for occupancy function with motion-detection data...')

        path = os.getcwd() #Get current directory
        output_path = path + r'\reports\6-occupancy' #Specify the output path for report
        input_path = path + r'\outputs\6-occupancy' #Specify the input path for report

        #Extract KPIs excel sheet
        kpis = pd.read_excel(input_path + r'\motion_detection_kpis.xlsx',sheet_name='KPIs',keep_default_na=False)
        kpis.drop(kpis.columns[0],axis=1,inplace=True)

        #Generate Word document
        document = Document()

        #Report title
        document.add_heading('Occupancy (Motion-detection) - Analysis Report', 0)

        #Introductory paragraph
        p = document.add_paragraph('The occupancy function determines the typical ')
        p.add_run('earliest arrival time, latest arrival time, latest departure time,').bold = True
        p.add_run(' and ')
        p.add_run('longest break duration').bold = True
        p.add_run(' for weekdays ONLY. This function can be used to inform ventilation schedules which can minimize \
    excessive ventilation during unoccupied hours, or even serve as a bssis for an occupant-driven demand controlled ventilation scheme. \
    Visuals plot building-level and floor-level occupant count profiles.')

        #KPIs heading and description
        document.add_heading('Key performance Indicators', level=1)
        p = document.add_paragraph('This section presents the generated KPIs - ')
        p.add_run('typical earliest arrival time, latest arrival time, latest departure time, and longest break duration').bold = True
        p.add_run(' for weekdays ONLY.')

        #Add tables with summary tables
        table_labels = ['Earlist arrival time',
                        'Latest arrival time',
                        'Latest departure time',
                        'Longest break duration (hours)'
                        ]

        t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
        
        for j in range(kpis.shape[-1]): # add the header rows.
            t.cell(0,j).text = table_labels[j]
        
        for i in range(kpis.shape[0]): # add the rest of the data frame
            for j in range(kpis.shape[-1]):
                t.cell(i+1,j).text = str(kpis.values[i,j])
        
        t.style = 'Colorful List'

        #Save document in reports folder
        document.save(output_path + r'\occupancy_motion_report.docx')
        convert(output_path + r'\occupancy_motion_report.docx', output_path + r'\occupancy_motion_report.pdf')
        print('Report successfully generated!')

        return

#------------------------------------------------------GENERATE COMPLAINTS REPORT-------------------------------------------------
def complaints():
    print('Generating report for occupant complaints analytics function...')

    path = os.getcwd() #Get current directory
    output_path = path + r'\reports\7-complaintsAnalytics' #Specify the output path for report
    input_path = path + r'\outputs\7-complaintsAnalytics' #Specify the input path for report

    #Extract KPIs excel sheet
    kpis = pd.read_excel(input_path + r'\complaints_freq.xlsx',sheet_name='Daily frequency of complaints')
    kpis.drop(kpis.columns[0],axis=1,inplace=True)
    kpis.iloc[:,1:] = kpis.iloc[:,1:].astype(float).round(3)

    #Generate Word document
    document = Document()

    #Report title
    document.add_heading('Complaint Analytics - Analysis Report', 0)

    #Introductory paragraph
    p = document.add_paragraph('The occupant complaints function determines the ')
    p.add_run('daily frequency of hot/cold-related occupant complaints. ').bold = True
    p.add_run('This function can be used to assess the effects of temperature setpoint or \
schedule adjustments on building occupants. The visuals depict the distribution of complaints by month and by the period \
of the day, the relationship between the complaints and indoor and outdoor air temperature, and the predicted proportion \
of the complaints based on time of day, day of the week, and outdoor air temperature. More information is available at the \
respective sections.')

    #First visualization heading and description
    document.add_heading('Visuals - Occupant complaints breakdown', level=1)
    p = document.add_paragraph('This set of visuals categorizes the complaints by the type of complaint \
(hot or cold related), and counts the number of complaints by the month and period of the day the complaint was \
registered.')

    #Add first visualizations
    document.add_picture(input_path + r'\complaints_breakdown.png', width=Inches(5.75))

    #Second visualization heading and description
    document.add_heading('Visuals - Indoor and outdoor air temperature', level=1)
    p = document.add_paragraph('This visual plots the relationship between a hot/cold complaint \
and the indoor and outdoor air temperature at the time the complaint was registered.')

    #Add second visualizations
    document.add_picture(input_path + r'\complaint_scatter.png', width=Inches(5.75))

    #Third visualization heading and description
    document.add_heading('Visuals - Predicted proportion of complaints', level=1)
    p = document.add_paragraph("This visual predicts the proportion of complaints that would be made \
with respect to certain conditions. The boxes which branch to the left represent the predicted proportion of complaints when \
the condition in the preceding box is satisfied. For example, if a box with the condition, 'Hour of the day <= 10' \
has a left branch with "+r"40%"+" and a right branch with "+r"60%"+", it is predicted that "+r"40%"+" of all complaints \
made will occur before 10 am. The condition is displayed at the top of the box, and the predicted \
proportion is displayed at the center of each box between the 0.0s.")

    #Add third visualizations
    document.add_picture(input_path + r'\decision_tree.png', width=Inches(5.75))

    #KPIs heading and description
    document.add_heading('Key performance Indicators', level=1)
    p = document.add_paragraph('This section presents the generated KPIs - ')
    p.add_run('daily frequency of hot and cold complaints in the heating and cooling season.').bold = True
    p.add_run(' The values represent the number of complaints made per day. Higher frequencies indicate a higher \
rate of occurence of a type of complaint for the particular season.')

    table_labels = ['Type of complaint',
                    'Daily frequency for heating season (complaints per day)',
                    'Daily frequency for cooling season (complaints per day)']

    p = document.add_heading('Daily frequencies of hot/cold occupant complaints',level=2)
    t = document.add_table(kpis.shape[0]+1, kpis.shape[1])
    
    for j in range(kpis.shape[-1]): # add the header rows.
        t.cell(0,j).text = table_labels[j]
    
    for i in range(kpis.shape[0]): # add the rest of the data frame
        for j in range(kpis.shape[-1]):
            t.cell(i+1,j).text = str(kpis.values[i,j])
    
    t.style = 'Colorful List'

    #Save document in reports folder
    document.save(output_path + r'\complaint_report.docx')
    convert(output_path + r'\complaint_report.docx', output_path + r'\complaint_report.pdf')
    print('Report successfully generated!')

    return